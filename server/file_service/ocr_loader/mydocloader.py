from typing import List

import tqdm
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader
from langchain_core.documents import Document
from server.file_service.ocr_loader.ocr import get_ocr

class RapidOCRDocLoader(UnstructuredFileLoader):

    """
    重写 _get_elements 方法，实现对 docx 文档的解析
    """
    def _get_elements(self) -> List:
 
        def doc2text(filepath):
            from io import BytesIO
            import numpy as np
            from docx import ImagePart
            from docx import Document as DocxDocument  # 重命名避免和 langchain 的 Document 冲突
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table, _Cell
            from docx.text.paragraph import Paragraph
            from PIL import Image

            # 拿到 rapidOcr 实例
            ocr = get_ocr()
            # 加载 doc 文档
            doc = DocxDocument(filepath)
            # 初始化空字符串用于拼接所有提取的文本呢
            resp = ""


            """
            递归遍历 doc 文档中的所有段落和表格，提取文本   
            """
            def iter_block_items(parent):
                from docx.document import Document as DocxDocument

                # 判断父节点类型，确定要遍历的根元素
                if isinstance(parent, DocxDocument):
                    # 如果是整个文档，遍历 body 下的所有子元素
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    # 如果是表格单元格，遍历单元格内的子元素
                    parent_elm = parent._tc
                else:
                    # 不支持的类型，抛出异常
                    raise ValueError("RapidOCRDocLoader parse fail")

                # 遍历父元素的所有子节点
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        # 如果是段落节点，返回 Paragraph 对象
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        # 如果是表格节点，返回 Table 对象
                        yield Table(child, parent)



            # 初始化进度条，显示当前处理的块索引和总块数的比值
            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="RapidOCRDocLoader block index: 0",
            )
            # 遍历 docx 中的每一个块（段落/表格）
            for i, block in enumerate(iter_block_items(doc)):
                # 更新进度条描述，显示当前处理的块索引
                b_unit.set_description("RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()

                # 如果是段落块，提取文字与图片（图片进行 ocr）
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    images = block._element.xpath(".//pic:pic")  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # 获取图片id
                            part = doc.part.related_parts[
                                img_id
                            ]  # 根据图片id获取对应的图片部分
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)
                
                # 如果是表格块：提取每个单元格的文字
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                
                # 进度条更新（处理完一个块，进度+1）
                b_unit.update(1)
            return resp

        # 调用 doc2text 函数提取文本
        text = doc2text(self.file_path)

        # 替代原有的 return partition_text(...)，无需分词，仅封装文本
        return [Document(page_content=text.strip(), metadata={"source": self.file_path})]
 
        # from unstructured.partition.text import partition_text
        # return partition_text(text=text, **self.unstructured_kwargs)


if __name__ == "__main__":
    from pathlib import Path
    file_path = Path(__file__).parent.parent.parent / "test" / "samples" / "ocr_test.docx"
    loader = RapidOCRDocLoader(file_path=str(file_path))
    docs = loader.load()
    print(docs)